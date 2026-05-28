# core/export_tools.py

import os
import json
import logging

logger = logging.getLogger(__name__)

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

def generate_py(raw_content: str, output_path: str) -> str:
    """Generates a Python (.py) file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(raw_content)
    return output_path

def generate_md(raw_content: str, output_path: str) -> str:
    """Generates a Markdown (.md) file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(raw_content)
    return output_path

def generate_txt(raw_content: str, output_path: str) -> str:
    """Generates a Plain Text (.txt) file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(raw_content)
    return output_path

def generate_json(raw_content: str, output_path: str) -> str:
    """Generates a JSON (.json) file."""
    try:
        # Validate JSON format before saving
        data = json.loads(raw_content)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)
    except json.JSONDecodeError as e:
        logger.warning(f"Failed to parse JSON, saving as raw text instead: {e}")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(raw_content)
    return output_path

def generate_docx(raw_content: str, output_path: str) -> str:
    """Generates a Word Document (.docx) file."""
    if Document is None:
        logger.error("python-docx is not installed. Falling back to Markdown.")
        fallback_path = output_path.replace(".docx", ".md")
        return generate_md(raw_content, fallback_path)

    doc = Document()
    
    # Simple markdown-like parser to add headers and paragraphs
    lines = raw_content.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            doc.add_paragraph(stripped)
            
    doc.save(output_path)
    return output_path

def generate_xlsx(raw_content: str, output_path: str) -> str:
    """Generates an Excel Document (.xlsx) file. Expects JSON array or CSV format."""
    if pd is None:
        logger.error("pandas is not installed. Falling back to text.")
        fallback_path = output_path.replace(".xlsx", ".txt")
        return generate_txt(raw_content, fallback_path)

    try:
        # Try to parse as JSON first (array of dicts)
        data = json.loads(raw_content)
        df = pd.DataFrame(data)
        df.to_excel(output_path, index=False)
        return output_path
    except Exception:
        # If it's not JSON, try to treat it as raw CSV string
        import io
        try:
            df = pd.read_csv(io.StringIO(raw_content))
            df.to_excel(output_path, index=False)
            return output_path
        except Exception as e:
            logger.error(f"Could not parse content into Excel: {e}. Falling back to text.")
            fallback_path = output_path.replace(".xlsx", ".txt")
            return generate_txt(raw_content, fallback_path)

def generate_email(raw_content: str, output_path: str) -> str:
    """
    Generates an email export. The LLM is expected to produce content in the format:
    Subject: <subject line>
    To: <recipient(s)>
    ---
    <body text>
    
    If the format is not matched, the entire content is used as the body.
    The email content is saved to a .txt file for archival.
    Actual sending is attempted via the email tool if credentials are configured.
    """
    import os as _os

    # Parse subject/to/body from LLM output
    subject = "Alfredo Workflow Report"
    to = ""
    body = raw_content

    lines = raw_content.split('\n')
    body_start = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.lower().startswith("subject:"):
            subject = stripped[len("subject:"):].strip()
            body_start = i + 1
        elif stripped.lower().startswith("to:"):
            to = stripped[len("to:"):].strip()
            body_start = i + 1
        elif stripped == "---":
            body_start = i + 1
            break

    body = '\n'.join(lines[body_start:]).strip() if body_start > 0 else raw_content

    # Save email as .txt for archival
    txt_path = output_path.replace(".email", ".txt") if output_path.endswith(".email") else output_path
    if not txt_path.endswith(".txt"):
        txt_path += ".txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(f"Subject: {subject}\nTo: {to}\n---\n{body}")

    # Attempt to send via email tool if credentials are available
    try:
        if to and _os.getenv("OUTLOOK_EMAIL"):
            import tools.email_tool as email_tool
            result = email_tool.send_email(to=to, subject=subject, body=body)
            logger.info(f"Email export sent: {result}")
    except Exception as e:
        logger.warning(f"Could not send email export (saved locally): {e}")

    return txt_path


# Map string keys to specific generator functions and their default extensions
EXPORT_TOOL_MAP = {
    "python": {"func": generate_py, "ext": "py"},
    "markdown": {"func": generate_md, "ext": "md"},
    "text": {"func": generate_txt, "ext": "txt"},
    "json": {"func": generate_json, "ext": "json"},
    "word": {"func": generate_docx, "ext": "docx"},
    "excel": {"func": generate_xlsx, "ext": "xlsx"},
    "email": {"func": generate_email, "ext": "txt"},
}

