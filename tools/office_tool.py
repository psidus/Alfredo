"""
tools/office_tool.py

Agent tools for creating and modifying Microsoft Office documents (Word, Excel)
and taking screenshots.

SECURITY MODEL:
  - All WRITE operations (create/edit documents) require explicit human confirmation
    via Telegram before they are executed.
  - The agent sends a preview of what it intends to do, then blocks until the operator
    replies CONFERMA or ANNULLA.
  - READ operations (listing content) do not require confirmation.
  - Screenshots are taken and sent directly to the operator without requiring confirmation
    (they are non-destructive).
"""

import os
import logging
import io
from crewai.tools import tool

logger = logging.getLogger(__name__)


def _request_confirmation(preview_message: str) -> bool:
    """
    Sends a confirmation request to the Telegram operator and blocks until a response.
    Returns True if confirmed, False if cancelled or timeout.
    """
    from core.human_in_the_loop import request_human_input
    from core.notification_manager import NotificationManager

    chat_id = os.environ.get("CURRENT_CHAT_ID")
    if not chat_id:
        logger.warning("No CURRENT_CHAT_ID set — auto-denying write operation.")
        return False

    full_message = (
        "⚠️ <b>CONFIRMATION REQUIRED</b>\n\n"
        f"{preview_message}\n\n"
        "Reply:\n"
        "✅ <code>CONFIRM</code> — to proceed\n"
        "❌ <code>CANCEL</code> — to abort"
    )

    notifier = NotificationManager()
    notifier.send_telegram_notification(full_message, chat_id=chat_id)

    answer = request_human_input(chat_id, full_message)
    answer_clean = answer.strip().upper() if answer else ""

    if answer_clean in ("CONFIRM", "CONFERMA"):  # Accept both languages
        logger.info("Operator confirmed write operation.")
        return True
    else:
        logger.info(f"Operator denied write operation (replied: '{answer}').")
        return False


def _send_file_to_telegram(file_path: str, caption: str = "") -> bool:
    """Sends a file to the Telegram operator."""
    try:
        from core.notification_manager import NotificationManager
        chat_id = os.environ.get("CURRENT_CHAT_ID")
        if not chat_id:
            return False

        notifier = NotificationManager()
        notifier.send_file_notification(file_path, caption=caption, chat_id=chat_id)
        return True
    except Exception as e:
        logger.error(f"Failed to send file to Telegram: {e}")
        return False


# ===========================================================================
# WORD DOCUMENT TOOLS
# ===========================================================================

@tool
def create_word_document(file_path: str, title: str, content: str) -> str:
    """
    Creates a new Microsoft Word (.docx) document at the specified full path.
    The document will have a title and body content.

    IMPORTANT: Before creating the file, this tool will ask the operator for confirmation
    via Telegram, showing a preview of the document content.

    Args:
        file_path: Full absolute path where the document should be saved (e.g., 'C:/Users/User/Documents/report.docx').
        title: The document title (will be a Heading 1).
        content: The body text of the document. Use '\\n\\n' to separate paragraphs.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return "Error: 'python-docx' library not installed. Run: pip install python-docx"

    abs_path = os.path.abspath(file_path)
    workspace_dir = os.path.abspath("workspace")
    if not abs_path.startswith(workspace_dir):
        return f"Error: Security violation. Path '{file_path}' is outside the designated workspace/ directory."

    # Build preview
    content_preview = content[:500] + ("..." if len(content) > 500 else "")
    preview = (
        f"📄 <b>Create Word Document</b>\n"
        f"📍 Path: <code>{abs_path}</code>\n"
        f"📋 Title: <b>{title}</b>\n"
        f"📝 Content preview:\n<pre>{content_preview}</pre>"
    )

    confirmed = _request_confirmation(preview)
    if not confirmed:
        return "Operation cancelled by the operator. No file was created."

    try:
        doc = Document()

        # Title
        heading = doc.add_heading(title, level=1)

        # Body — split by double newline as paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())

        # Ensure directory exists
        os.makedirs(os.path.dirname(abs_path), exist_ok=True)
        doc.save(abs_path)

        # Send file to Telegram as confirmation
        _send_file_to_telegram(abs_path, caption=f"✅ Word document created: {os.path.basename(abs_path)}")

        return f"✅ Word document created successfully: '{abs_path}'"

    except Exception as e:
        logger.error(f"Error creating Word document at '{abs_path}': {e}", exc_info=True)
        return f"Error creating Word document: {e}"


@tool
def edit_word_document(file_path: str, new_content: str, append: bool = False) -> str:
    """
    Edits an existing Microsoft Word (.docx) document by replacing its body content
    or appending new paragraphs to it.

    IMPORTANT: Before modifying the file, this tool asks the operator for confirmation
    via Telegram, showing what will change.

    Args:
        file_path: Full absolute path to the existing .docx file.
        new_content: New content to write. Use '\\n\\n' to separate paragraphs.
        append: If True, adds content to the end. If False (default), replaces all body text.
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: 'python-docx' library not installed. Run: pip install python-docx"

    abs_path = os.path.abspath(file_path)
    workspace_dir = os.path.abspath("workspace")
    if not abs_path.startswith(workspace_dir):
        return f"Error: Security violation. Path '{file_path}' is outside the designated workspace/ directory."
    if not os.path.isfile(abs_path):
        return f"Error: File not found at '{abs_path}'."

    action_label = "APPEND (add to end)" if append else "FULL REPLACEMENT of existing content"
    content_preview = new_content[:400] + ("..." if len(new_content) > 400 else "")
    preview = (
        f"✏️ <b>Edit Word Document</b>\n"
        f"📍 Path: <code>{abs_path}</code>\n"
        f"🔄 Action: {action_label}\n"
        f"📝 New content:\n<pre>{content_preview}</pre>"
    )

    confirmed = _request_confirmation(preview)
    if not confirmed:
        return "Operation cancelled by the operator. The file was not modified."

    try:
        doc = Document(abs_path)

        if not append:
            # Clear all existing paragraphs
            for para in doc.paragraphs:
                p = para._element
                p.getparent().remove(p)

        paragraphs = new_content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

        doc.save(abs_path)

        _send_file_to_telegram(abs_path, caption=f"✅ Word document edited: {os.path.basename(abs_path)}")
        return f"✅ Word document edited successfully: '{abs_path}'"

    except Exception as e:
        logger.error(f"Error editing Word document at '{abs_path}': {e}", exc_info=True)
        return f"Error editing Word document: {e}"


# ===========================================================================
# EXCEL DOCUMENT TOOLS
# ===========================================================================

@tool
def create_excel_document(file_path: str, sheet_name: str, headers: str, rows: str) -> str:
    """
    Creates a new Microsoft Excel (.xlsx) document at the specified full path.

    IMPORTANT: Before creating the file, this tool asks the operator for confirmation.

    Args:
        file_path: Full absolute path where the .xlsx should be saved.
        sheet_name: Name of the worksheet (e.g., 'Sales Report').
        headers: Comma-separated column headers (e.g., 'Name,Age,City').
        rows: Each row as a semicolon-separated line, with values comma-separated.
              Example: 'Alice,30,Rome;Bob,25,Milan;Charlie,35,Turin'
    """
    try:
        import openpyxl
    except ImportError:
        return "Error: 'openpyxl' library not installed. Run: pip install openpyxl"

    abs_path = os.path.abspath(file_path)
    workspace_dir = os.path.abspath("workspace")
    if not abs_path.startswith(workspace_dir):
        return f"Error: Security violation. Path '{file_path}' is outside the designated workspace/ directory."

    # Parse headers and rows for preview
    header_list = [h.strip() for h in headers.split(',')]
    row_list = []
    for raw_row in rows.strip().split(';'):
        if raw_row.strip():
            row_list.append([v.strip() for v in raw_row.split(',')])

    preview_rows = row_list[:5]
    preview_table = " | ".join(header_list) + "\n"
    preview_table += "-" * (len(preview_rows[0]) * 12 if preview_rows else 40) + "\n"
    for r in preview_rows:
        preview_table += " | ".join(r) + "\n"
    if len(row_list) > 5:
        preview_table += f"... and {len(row_list) - 5} more rows"

    preview = (
        f"📊 <b>Create Excel File</b>\n"
        f"📍 Path: <code>{abs_path}</code>\n"
        f"📋 Sheet: <b>{sheet_name}</b>\n"
        f"🔢 Rows: {len(row_list)}\n"
        f"<pre>{preview_table}</pre>"
    )

    confirmed = _request_confirmation(preview)
    if not confirmed:
        return "Operation cancelled by the operator. No file was created."

    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Write headers in bold
        from openpyxl.styles import Font
        ws.append(header_list)
        for cell in ws[1]:
            cell.font = Font(bold=True)

        # Write data rows
        for row in row_list:
            ws.append(row)

        # Auto-width columns
        for column_cells in ws.columns:
            max_length = max((len(str(cell.value or "")) for cell in column_cells), default=10)
            ws.column_dimensions[column_cells[0].column_letter].width = min(max_length + 4, 40)

        os.makedirs(os.path.dirname(abs_path), exist_ok=True)
        wb.save(abs_path)

        _send_file_to_telegram(abs_path, caption=f"✅ Excel file created: {os.path.basename(abs_path)}")
        return f"✅ Excel file created successfully: '{abs_path}' ({len(row_list)} rows)"

    except Exception as e:
        logger.error(f"Error creating Excel file at '{abs_path}': {e}", exc_info=True)
        return f"Error creating Excel document: {e}"


# ===========================================================================
# SCREENSHOT TOOL
# ===========================================================================

@tool
def take_screenshot(note: str = "") -> str:
    """
    Takes a screenshot of the current screen and sends it to the operator via Telegram.
    Use this to show the operator the current state of the screen,
    or to document what is happening on the computer.

    Args:
        note: An optional note to include with the screenshot (e.g., 'This is the error I found').
    """
    try:
        import mss
        import mss.tools
        from PIL import Image
    except ImportError:
        return "Error: 'mss' or 'Pillow' not installed. Run: pip install mss Pillow"

    try:
        screenshot_path = os.path.abspath("workspace/_screenshot_temp.png")
        os.makedirs(os.path.dirname(screenshot_path), exist_ok=True)

        with mss.mss() as sct:
            # Capture all monitors
            monitor = sct.monitors[0]  # 0 = all monitors combined
            sct_img = sct.grab(monitor)
            mss.tools.to_png(sct_img.rgb, sct_img.size, output=screenshot_path)

        # Resize if too large (max 1920x1080 for Telegram)
        img = Image.open(screenshot_path)
        img.thumbnail((1920, 1080), Image.LANCZOS)
        img.save(screenshot_path, optimize=True, quality=85)

        caption = "📸 Screenshot from operator's machine"
        if note:
            caption += f"\n📝 Agent note: {note}"

        sent = _send_file_to_telegram(screenshot_path, caption=caption)

        if sent:
            return f"✅ Screenshot captured and sent to the operator via Telegram.{(' Note: ' + note) if note else ''}"
        else:
            return "Screenshot captured but could not be sent (CURRENT_CHAT_ID not set)."

    except Exception as e:
        logger.error(f"Error taking screenshot: {e}", exc_info=True)
        return f"Error taking screenshot: {e}"
